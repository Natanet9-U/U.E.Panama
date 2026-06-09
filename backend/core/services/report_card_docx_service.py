import io

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..tracing import trace_service_class
from .access_service import AccessControlService
from .report_card_service import ReportCardService


AREA_MAP = {
    'Lenguaje y Comunicacion': 'Comunidad y Sociedad',
    'Ciencias Sociales': 'Comunidad y Sociedad',
    'Educacion Musical': 'Comunidad y Sociedad',
    'Artes Plasticas': 'Comunidad y Sociedad',
    'Educacion Fisica': 'Comunidad y Sociedad',
    'Matematicas': 'Ciencia, Tecnolog\u00eda y Producci\u00f3n',
    'Tecnica Tecnologica': 'Ciencia, Tecnolog\u00eda y Producci\u00f3n',
    'Ciencias Naturales': 'Vida, Tierra y Territorio',
    'Valores,Ecspiritualidad y Religiones': 'Cosmos y Pensamiento',
}

AREA_ORDER = [
    'Lenguaje y Comunicacion',
    'Ciencias Sociales',
    'Educacion Musical',
    'Artes Plasticas',
    'Educacion Fisica',
    'Matematicas',
    'Tecnica Tecnologica',
    'Ciencias Naturales',
    'Valores,Ecspiritualidad y Religiones',
]

_UNIDADES = ['CERO', 'UNO', 'DOS', 'TRES', 'CUATRO', 'CINCO', 'SEIS', 'SIETE', 'OCHO', 'NUEVE']
_DECENAS = ['', 'DIEZ', 'VEINTE', 'TREINTA', 'CUARENTA', 'CINCUENTA',
            'SESENTA', 'SETENTA', 'OCHENTA', 'NOVENTA']
_ESPECIALES = {11: 'ONCE', 12: 'DOCE', 13: 'TRECE', 14: 'CATORCE',
               15: 'QUINCE', 16: 'DIECISEIS', 17: 'DIECISIETE',
               18: 'DIECIOCHO', 19: 'DIECINUEVE',
               21: 'VEINTIUNO', 22: 'VEINTIDOS', 23: 'VEINTITRES',
               24: 'VEINTICUATRO', 25: 'VEINTICINCO', 26: 'VEINTISEIS',
               27: 'VEINTISIETE', 28: 'VEINTIOCHO', 29: 'VEINTINUEVE'}


def _numero_a_texto(n):
    if n in _ESPECIALES:
        return _ESPECIALES[n]
    if n < 10:
        return _UNIDADES[n]
    if n < 100:
        d = n // 10
        u = n % 10
        if u == 0:
            return _DECENAS[d]
        if d == 2:
            return f'VEINTI{_UNIDADES[u]}'
        return f'{_DECENAS[d]} Y {_UNIDADES[u]}'
    if n == 100:
        return 'CIEN'
    return str(n)


@trace_service_class
class ReportCardDOCXService:

    def __init__(self):
        self.ac = AccessControlService()
        self._rcs = ReportCardService()

    def generar_docx(self, usuario, estudiante_id, gestion=None):
        data = self._rcs.generar_boletin(usuario, estudiante_id, gestion)
        return self._build_docx(data)

    def _build_docx(self, data):
        doc = Document('/app/templates/boletin_template.docx')

        est = data['estudiante']
        curso = data['curso']
        periodos = data['periodos']

        materias_por_area = {m['area']: m for m in data['materias']}
        materias_ordenadas = []
        for area in AREA_ORDER:
            if area in materias_por_area:
                materias_ordenadas.append(materias_por_area.pop(area))
        for m in data['materias']:
            if m['area'] in materias_por_area:
                materias_ordenadas.append(materias_por_area.pop(m['area']))

        # --- Table 2 (index 2): Grades ---
        t = doc.tables[2]

        # Row 0: student info
        row0 = t.rows[0]
        self._set_cell_text(row0.cells[0], 'C\u00f3digo Rude: ', est['rude'])
        apellidos = f"{est['primer_apellido']} {est['segundo_apellido']}".strip()
        self._set_cell_text(row0.cells[2], 'Apellido y Nombres: ', f'{apellidos} {est["nombres"]}')
        self._set_cell_text(row0.cells[6], 'A\u00f1o de Escolaridad: ', f'{curso["grado"]} {curso["paralelo"]}')
        for p in row0.cells[6].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for i, mat in enumerate(materias_ordenadas):
            row_idx = 5 + i
            if row_idx >= len(t.rows):
                self._add_subject_row(t, row_idx)
            row = t.rows[row_idx]

            campo = AREA_MAP.get(mat['area'], '')
            self._set_campo_cell(row.cells[0], campo)
            row.cells[1].text = mat['area']

            for pi, p in enumerate(periodos):
                val = mat['notas_por_periodo'].get(str(p['id']))
                col_idx = 3 + pi
                if col_idx < len(row.cells):
                    redondeado = round(val) if val is not None else None
                    self._set_centered_cell(row.cells[col_idx], str(redondeado) if redondeado is not None else '-')

            prom = mat['promedio_final']
            prom_redondeado = round(prom) if prom is not None else None
            self._set_centered_cell(row.cells[7], str(prom_redondeado) if prom_redondeado is not None else '-')
            self._set_centered_cell(row.cells[8], _numero_a_texto(prom_redondeado) if prom_redondeado is not None else '-')

        for i in range(len(materias_ordenadas), len(t.rows) - 5):
            row_idx = 5 + i
            if row_idx < len(t.rows):
                for cell in t.rows[row_idx].cells:
                    cell.text = ''

        # --- Table 3 (index 3): Promotion result ---
        t3 = doc.tables[3]
        notas = [m['promedio_final'] for m in materias_ordenadas if m['promedio_final'] is not None]
        promedio_general = round(sum(notas) / len(notas)) if notas else None
        estado = 'APROBADO' if promedio_general is not None and promedio_general >= 51 else 'REPROBADO'
        self._set_promocion_cell(t3.rows[0].cells[0], estado)

        # --- Apply Times New Roman 7pt to tables 2 and 3 (grades + promotion) ---
        self._set_font(doc, 'Times New Roman', 7, table_indices={2, 3})

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_subject_row(self, table, after_index):
        from copy import deepcopy
        ref_row = table.rows[after_index - 1]
        tr = deepcopy(ref_row._tr)
        table._tbl.append(tr)

    def _set_cell_text(self, cell, label, value):
        cell.text = ''
        p = cell.paragraphs[0]
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(7)
        run_val = p.add_run(str(value))
        run_val.font.name = 'Times New Roman'
        run_val.font.size = Pt(7)

    def _set_campo_cell(self, cell, texto):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(7)

    def _set_centered_cell(self, cell, texto):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(7)

    def _set_promocion_cell(self, cell, estado):
        cell.text = ''
        p = cell.paragraphs[0]
        run_label = p.add_run('Informe de Promoci\u00f3n: ')
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(7)
        run_val = p.add_run(estado)
        run_val.bold = True
        run_val.font.name = 'Times New Roman'
        run_val.font.size = Pt(7)

    @staticmethod
    def _set_font(doc, font_name, font_size, table_indices=None):
        for p in doc.paragraphs:
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
        for idx, t in enumerate(doc.tables):
            if table_indices is not None and idx not in table_indices:
                continue
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)

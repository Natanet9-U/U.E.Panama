from io import BytesIO

from django.db import connection
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, Side, PatternFill
from docx import Document
from docx.shared import Inches

from django.db.models import Count, Q, Value, TextField
from django.db.models.functions import Concat, Coalesce, Trim

from ..models import DocenteAsignacion, ExportEvent, Inscripciones, Periodos
from .access_service import AccessControlService
from ..tracing import trace_service_class
from ..models import AuditLog
from django.db.models import Count
from django.utils.dateparse import parse_date


@trace_service_class
class ReportsService:

    def __init__(self):
        self.ac = AccessControlService()

    # ── Excel Export ──────────────────────────────────────────────────────────────

    def export_notas_excel(self, usuario, docente_asignacion_id, periodo_id):
        if not self.ac.puede_exportar(usuario):
            raise PermissionError('No tienes permisos para exportar')

        da = DocenteAsignacion.objects.select_related(
            'curso__grado__nivel', 'curso__paralelo', 'area', 'docente__usuario'
        ).get(id=docente_asignacion_id)

        periodo = Periodos.objects.get(id=periodo_id)

        estudiantes = Inscripciones.objects.filter(
            curso=da.curso, gestion=da.gestion, estado='activo'
        ).select_related('estudiante')

        notas_dim = self._get_notas_dim(da.id, periodo.id)
        notas_tot = self._get_notas_tot(da.id, periodo.id)

        wb = Workbook()
        ws = wb.active
        ws.title = 'Notas'

        # Header
        ws.cell(row=1, column=1, value=f'{da.area.nombre} - {da.curso}')
        ws.merge_cells('A1:H1')
        ws.cell(row=2, column=1, value=f'{periodo.nombre} {periodo.gestion}')
        ws.merge_cells('A2:H2')
        ws.cell(row=3, column=1, value=f'Docente: {da.usuario.nombre_completo}')

        # Column headers
        headers = ['N°', 'RUDE', 'CI', 'Nombres', 'Apellido', 'SER', 'SABER', 'HACER', 'AUTOEV', 'TOTAL']
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=h)
            cell.font = Font(bold=True)
            cell.border = Border(
                bottom=Side(style='thin'),
                top=Side(style='thin'),
                left=Side(style='thin'),
                right=Side(style='thin'),
            )

        # Data
        for i, ins in enumerate(estudiantes, 1):
            e = ins.estudiante
            row = i + 5
            ws.cell(row=row, column=1, value=i)
            ws.cell(row=row, column=2, value=e.rude)
            ws.cell(row=row, column=3, value=e.ci)
            ws.cell(row=row, column=4, value=e.nombres)
            ws.cell(row=row, column=5, value=e.primer_apellido)

            est_notas = notas_dim.get(e.id, {})
            total = notas_tot.get(e.id, 0)

            ws.cell(row=row, column=6, value=est_notas.get('SER'))
            ws.cell(row=row, column=7, value=est_notas.get('SABER'))
            ws.cell(row=row, column=8, value=est_notas.get('HACER'))
            ws.cell(row=row, column=9, value=est_notas.get('AUTOEVALUACION'))
            ws.cell(row=row, column=10, value=total)

            for col in range(1, 11):
                ws.cell(row=row, column=col).border = Border(
                    bottom=Side(style='thin'),
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                )

        # Column widths
        for col, w in [(1, 5), (2, 14), (3, 12), (4, 25), (5, 25), (6, 8), (7, 8), (8, 8), (9, 10), (10, 8)]:
            ws.column_dimensions[chr(64 + col)].width = w

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf

    def _get_notas_dim(self, docente_asignacion_id, periodo_id):
        with connection.cursor() as cursor:
            cursor.execute(
                """SELECT v.estudiante_id, d.nombre, v.nota_dimension
                   FROM v_notas_por_dimension v
                   JOIN dimensiones_evaluacion d ON d.id = v.dimension_id
                   WHERE v.docente_asignacion_id = %s AND v.periodo_id = %s""",
                [docente_asignacion_id, periodo_id],
            )
            rows = cursor.fetchall()
        result = {}
        for est_id, dim_nombre, nota in rows:
            result.setdefault(est_id, {})[dim_nombre] = float(nota) if nota else 0
        return result

    def _get_notas_tot(self, docente_asignacion_id, periodo_id):
        with connection.cursor() as cursor:
            cursor.execute(
                """SELECT estudiante_id, nota_total
                   FROM v_notas_totales
                   WHERE docente_asignacion_id = %s AND periodo_id = %s""",
                [docente_asignacion_id, periodo_id],
            )
            rows = cursor.fetchall()
        return {r[0]: float(r[1]) if r[1] else 0 for r in rows}

    def get_audit_load_summary(self, usuario, since=None, until=None, tabla=None):
        if not self.ac.puede_ver_auditoria(usuario):
            raise PermissionError('No tienes permisos para ver auditoria')

        qs = AuditLog.objects.all()
        if tabla:
            qs = qs.filter(tabla=tabla)
        if since:
            try:
                sd = parse_date(since)
                qs = qs.filter(fecha_cambio__date__gte=sd)
            except Exception:
                pass
        if until:
            try:
                ud = parse_date(until)
                qs = qs.filter(fecha_cambio__date__lte=ud)
            except Exception:
                pass

        agg = qs.values('usuario__id').annotate(
            changes=Count('id'),
            **{'usuario__nombre_completo': Trim(Concat(
                Coalesce('usuario__nombre', Value('', output_field=TextField())),
                Value(' ', output_field=TextField()),
                Coalesce('usuario__primer_apellido', Value('', output_field=TextField())),
                Value(' ', output_field=TextField()),
                Coalesce('usuario__segundo_apellido', Value('', output_field=TextField())),
                output_field=TextField(),
            ))}
        ).order_by('-changes')
        total = qs.count()
        return {'total_changes': total, 'by_user': list(agg)}

    def get_export_history(self, usuario, periodo_id=None, limit=20):
        if not self.ac.puede_exportar(usuario) and not self.ac.puede_ver_auditoria(usuario):
            raise PermissionError('No tienes permisos para ver exportaciones')

        qs = ExportEvent.objects.select_related('usuario', 'periodo').order_by('-creado_en')
        if periodo_id:
            qs = qs.filter(periodo_id=periodo_id)

        # Secretary/director can see all, teachers only their own exports
        if not (self.ac.es_secretaria(usuario) or self.ac.es_director(usuario)):
            qs = qs.filter(usuario=usuario)

        exports = []
        for entry in qs[:limit]:
            exports.append({
                'id': entry.id,
                'formato': entry.formato,
                'periodo_id': entry.periodo_id,
                'periodo': entry.periodo.nombre if entry.periodo else None,
                'gestion': entry.periodo.gestion if entry.periodo else None,
                'docente_asignacion_id': entry.docente_asignacion_id,
                'usuario': entry.usuario.nombre_completo if entry.usuario else None,
                'creado_en': entry.creado_en.isoformat(),
                'filtros': entry.filtros,
            })
        return {'exports': exports, 'total': qs.count()}

    def export_notas_docx(self, usuario, docente_asignacion_id, periodo_id):
        if not self.ac.puede_exportar(usuario):
            raise PermissionError('No tienes permisos para exportar')

        da = DocenteAsignacion.objects.select_related(
            'curso__grado__nivel', 'curso__paralelo', 'area', 'docente__usuario'
        ).get(id=docente_asignacion_id)

        periodo = Periodos.objects.get(id=periodo_id)

        estudiantes = Inscripciones.objects.filter(
            curso=da.curso, gestion=da.gestion, estado='activo'
        ).select_related('estudiante')

        notas_dim = self._get_notas_dim(da.id, periodo.id)
        notas_tot = self._get_notas_tot(da.id, periodo.id)

        doc = Document()
        doc.add_heading(f'{da.area.nombre} - {da.curso}', level=1)
        doc.add_paragraph(f'{periodo.nombre} {periodo.gestion}')
        doc.add_paragraph(f'Docente: {da.usuario.nombre_completo}')

        # Table headers
        headers = ['N°', 'RUDE', 'CI', 'Nombres', 'Apellido', 'SER', 'SABER', 'HACER', 'AUTOEV', 'TOTAL']
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for i, ins in enumerate(estudiantes, 1):
            e = ins.estudiante
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(e.rude or '')
            row_cells[2].text = str(e.ci or '')
            row_cells[3].text = str(e.nombres or '')
            row_cells[4].text = str(e.primer_apellido or '')

            est_notas = notas_dim.get(e.id, {})
            total = notas_tot.get(e.id, 0)

            row_cells[5].text = str(est_notas.get('SER') or '')
            row_cells[6].text = str(est_notas.get('SABER') or '')
            row_cells[7].text = str(est_notas.get('HACER') or '')
            row_cells[8].text = str(est_notas.get('AUTOEVALUACION') or '')
            row_cells[9].text = str(total)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
